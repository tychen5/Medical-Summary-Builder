from __future__ import annotations

from pathlib import Path
import datetime

from docx import Document as DocxDocument
from docx.table import Table

from ..schemas import MedicalSummary
from ..utils import ensure_directory


class ReportWriter:
    """Render structured medical summaries into DOCX and markdown outputs."""

    def __init__(self, output_dir: Path | str) -> None:
        self.output_dir = ensure_directory(output_dir)

    def write_markdown(self, summary: MedicalSummary, filename: str = "medical_summary.md") -> Path:
        """Persist the medical summary to markdown for auditing."""

        path = Path(self.output_dir) / filename
        content_lines: list[str] = [
            "# Medical Summary",
            "",
            "## Claimant Profile",
        ]
        for field_name, value in summary.profile.model_dump().items():
            content_lines.append(f"- **{field_name.replace('_', ' ').title()}**: {value or 'N/A'}")
        content_lines.append("")
        content_lines.append("## Timeline of Medical Events")
        for event in summary.events:
            content_lines.append(
                f"- {event.date or 'N/A'} | {event.provider or 'N/A'} | {event.reason or 'N/A'} | {event.reference or 'N/A'}"
            )
        path.write_text("\n".join(content_lines), encoding="utf-8")
        return path

    def write_docx(self, summary: MedicalSummary, template_path: Path | str | None = None) -> Path:
        """Persist the medical summary to a DOCX template."""

        document =  DocxDocument() #DocxDocument(template_path) if template_path else

        document.add_heading("Medical Summary", level=1)

        name = summary.profile.claimant_name or ""
        ssn = summary.profile.ssn or ""
        claim_title = summary.profile.claim_title or ""
        dli = self._format_date(summary.profile.date_last_insured)

        aod = self._format_date(summary.profile.alleged_onset_date)
        dob = self._format_date(summary.profile.date_of_birth)
        age_at_aod = self._format_integer(summary.profile.age_at_aod)
        current_age = self._format_integer(summary.profile.current_age)

        education = summary.profile.education or ""
        special_ed_status, special_ed_detail = self._extract_special_ed_info(summary)

        para_top = document.add_paragraph(
            f"NAME: {name}     SSN: {ssn}     Claim Title (II/XVI): {claim_title}     DLI: {dli}"
        )
        para_mid = document.add_paragraph(
            "AOD: "
            f"{aod}     Date of Birth: {dob}     Age at the time of AOD: {age_at_aod[:2]}     Current Age: {current_age[:2]}"
        )
        document.add_paragraph(
            "Last Grade Completed: "
            f"{education}     Attended Special Ed Classes: {special_ed_status}     Other Side Notes: {special_ed_detail}"#; (Ex {special_ed_detail})
        )

        document.add_paragraph("")

        table = document.add_table(rows=1, cols=4)
        self._apply_table_style(table, document, preferred_style="Light List")
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "DATE"
        hdr_cells[1].text = "PROVIDER"
        hdr_cells[2].text = "REASON"
        hdr_cells[3].text = "REF"

        for event in summary.events:
            row_cells = table.add_row().cells
            row_cells[0].text = self._format_date(event.date)
            row_cells[1].text = event.provider or ""
            row_cells[2].text = event.reason or ""
            row_cells[3].text = event.reference or ""

        output_path = Path(self.output_dir) / "medical_summary.docx"
        document.save(str(output_path))
        return output_path

    @staticmethod
    def _format_date(value) -> str:
        if isinstance(value, datetime.date):
            return value.strftime("%m/%d/%Y")
        if value is None:
            return ""
        # Handle datetime strings that may already be formatted
        try:
            parsed = datetime.datetime.fromisoformat(str(value)).date()
            return parsed.strftime("%m/%d/%Y")
        except ValueError:
            return str(value)

    @staticmethod
    def _format_integer(value) -> str:
        if value is None:
            return ""
        return str(value)

    def _extract_special_ed_info(self, summary: MedicalSummary) -> tuple[str, str]:
        status = "No"
        detail = "N/A"

        # Check notes for hints
        notes_lower = (summary.profile.notes or "").lower()
        if "special ed" in notes_lower or "special education" in notes_lower:
            status = "Yes"
            detail = summary.profile.notes or detail

        for table_name, rows in summary.custom_tables.items():
            if "education" not in table_name.lower():
                continue

            for row in rows:
                row_lower = {k.lower(): v for k, v in row.items()}
                if any("special" in key for key in row_lower):
                    raw_status = row_lower.get("special_ed") or row_lower.get("special education")
                    if isinstance(raw_status, str) and raw_status.strip():
                        status = raw_status.strip()
                    elif raw_status is True:
                        status = "Yes"
                    elif raw_status is False:
                        status = "No"

                    pieces: list[str] = []
                    for key, value in row.items():
                        if key.lower() in {"special_ed", "special education"}:
                            continue
                        if value:
                            pieces.append(str(value))
                    if pieces:
                        detail = ", ".join(pieces)
                    break

        return status or "No", detail or "N/A"

    @staticmethod
    def _apply_table_style(table: Table, document: DocxDocument, preferred_style: str) -> None:
        """Assign a table style if it exists in the template, otherwise fall back gracefully."""

        try:
            table.style = preferred_style
            return
        except (KeyError, ValueError):
            pass

        # Fall back to the default table style provided by the template/document.
        try:
            default_style = document.styles.table_styles.default
            if default_style is not None:
                table.style = default_style
        except (AttributeError, KeyError, ValueError):
            # As a last resort, leave the table unstyled.
            return
